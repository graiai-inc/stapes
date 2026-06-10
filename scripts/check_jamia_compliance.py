#!/usr/bin/env python3
"""Verify build/stapes_manuscript.docx against JAMIA Research and Applications
requirements (fetched from academic.oup.com/jamia/pages/General_Instructions
2026-06-10). Writes findings incrementally to build/jamia_compliance.txt.
"""

import re
from pathlib import Path

from docx import Document

BUILD = Path('/home/grey/dev/graiai/stapes/build')
OUT = BUILD / 'jamia_compliance.txt'

fh = open(OUT, 'w')


def emit(line: str = '') -> None:
    fh.write(line + '\n')
    fh.flush()
    print(line, flush=True)


doc = Document(str(BUILD / 'stapes_manuscript.docx'))
paras = doc.paragraphs

emit('=== JAMIA COMPLIANCE CHECK (Research and Applications) ===')
emit()

# --- Headings inventory ---
emit('[headings] all heading-style paragraphs in order:')
for p in paras:
    if p.style.name.startswith('Heading'):
        emit(f'   {p.style.name}: {p.text}')
emit()

# --- Title page block (everything before Abstract heading) ---
emit('[title-page] paragraphs before the Abstract heading:')
for p in paras:
    if p.style.name.startswith('Heading') and p.text.strip() == 'Abstract':
        break
    if p.text.strip():
        emit(f'   | {p.text}')
emit()

# --- Abstract structured labels ---
emit('[abstract] bold/labelled field names found in abstract block:')
in_abs = False
abs_text = []
for p in paras:
    t = p.text.strip()
    if p.style.name.startswith('Heading'):
        if t == 'Abstract':
            in_abs = True
            continue
        if in_abs:
            break
    if in_abs and t:
        abs_text.append(t)
        m = re.match(r'^([A-Za-z ]+):', t)
        if m:
            emit(f'   label: {m.group(1)}')
emit(f'   (abstract paragraphs: {len(abs_text)})')
emit()

# --- Tables ---
emit(f'[tables] docx contains {len(doc.tables)} tables (JAMIA limit: 4, Word format, placed where first cited)')
for i, tbl in enumerate(doc.tables, 1):
    emit(f'   table {i}: {len(tbl.rows)} rows x {len(tbl.columns)} cols')
emit()

# --- Figures (inline images) ---
img_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        img_count += 1
emit(f'[figures] inline images embedded in manuscript docx: {img_count} (JAMIA limit: 6; also uploaded separately)')
emit()

# --- Figure/table citation order in text ---
full_text = '\n'.join(p.text for p in paras)
fig_mentions = re.findall(r'Figure (\d+)', full_text)
tbl_mentions = re.findall(r'Table (\d+)', full_text)
emit(f'[citations] Figure mentions in order: {fig_mentions}')
emit(f'[citations] Table mentions in order: {tbl_mentions}')
emit()

# --- Required statements ---
emit('[statements] required-section presence:')
for needle in ['Background and Significance', 'Data Availability', 'Funding',
               'Conflict', 'Competing', 'Author Contributions', 'Acknowledg',
               'Code Availability', 'Ethics']:
    hits = [p.text.strip() for p in paras
            if p.style.name.startswith('Heading') and needle.lower() in p.text.lower()]
    emit(f'   {needle}: {hits if hits else "NOT FOUND as heading"}')
emit()

# --- Keywords ---
kw = [p.text for p in paras if 'keyword' in p.text.lower()]
emit(f'[keywords] lines mentioning keywords: {kw if kw else "NOT FOUND"}')
emit()

# --- Double spacing ---
from docx.shared import Length  # noqa: E402
spacings = set()
for p in paras[:60]:
    pf = p.paragraph_format
    spacings.add(pf.line_spacing)
emit(f'[spacing] distinct line_spacing values in first 60 paragraphs: {spacings}')
emit()

# --- Citation placement: [n] should be after punctuation ---
before_punct = re.findall(r'\[\d+(?:[,\-]\d+)*\][.,;]', full_text)
emit(f'[cite-style] citations appearing BEFORE punctuation (JAMIA wants after): {len(before_punct)}')
for ex in before_punct[:10]:
    emit(f'   example: ...{ex}')
emit()

emit('=== END ===')
fh.close()
