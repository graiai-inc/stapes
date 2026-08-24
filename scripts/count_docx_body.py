#!/usr/bin/env python3
"""Word-equivalent body count of the rendered submission docx.

Counts words in all paragraphs between the 'Background'
heading and the 'List of abbreviations' heading (the countable main text),
matching what an editor's Word count would see. Writes result to disk.
"""

import re
from pathlib import Path

from docx import Document

BUILD = Path('/home/grey/dev/graiai/stapes/build')
OUT = BUILD / 'qa_report.txt'

doc = Document(str(BUILD / 'stapes_manuscript.docx'))

in_body = False
words = 0
paras = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if text == 'Background':
        in_body = True
    if text == 'List of abbreviations':
        break
    # skip table captions: JAMIA word counts exclude tables, which now sit
    # inline where first cited (cell text is already outside doc.paragraphs)
    if re.match(r'^Table \d+\.', text):
        continue
    if in_body:
        words += len(p.text.split())
        paras += 1

with open(OUT, 'a') as fh:
    line = f'[docx-body] Word-equivalent body count (Background..Conclusion): {words} ({paras} paragraphs)'
    fh.write(line + '\n')
    fh.flush()
    print(line, flush=True)
