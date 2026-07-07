#!/usr/bin/env python3
"""Assemble the stapes npj Digital Medicine submission package.

Reads the individual paper components and produces:
    build/stapes_submission.md   — assembled manuscript markdown
    build/stapes_manuscript.docx — pandoc-converted Word file
    build/stapes_supplementary.md — supplementary tables markdown
    build/stapes_supplementary.docx — pandoc-converted Word file
    build/stapes_cover_letter.docx — cover letter Word file

Run with the lens venv python (it has pypandoc-binary installed):
    /home/grey/dev/graiai/lens/venv/bin/python scripts/assemble_submission.py
"""

import csv
import re
import subprocess
import sys
from pathlib import Path

import pypandoc

PAPER = Path('/home/grey/dev/graiai/stapes/paper')
BUILD = Path('/home/grey/dev/graiai/stapes/build')
BUILD.mkdir(parents=True, exist_ok=True)


def csv_to_md_table(csv_path: Path) -> str:
    """Render a CSV file as a GitHub-flavored markdown table."""
    with csv_path.open() as fh:
        reader = csv.reader(fh)
        rows = list(reader)
    if not rows:
        return ''
    header = rows[0]
    out = ['| ' + ' | '.join(header) + ' |']
    out.append('|' + '|'.join(['---'] * len(header)) + '|')
    for row in rows[1:]:
        out.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(out)


TITLE_PAGE = """# Privacy-Preserving On-Device Speech Recognition Achieves Near-Parity with Cloud Services for Clinical Documentation, but Both Fail on Medication Names

**Author:** J. Grey Faulkenberry, MD, MPH

Department of Hematology and Medical Oncology
Emory University School of Medicine
Atlanta, Georgia, USA

**Corresponding author:**
J. Grey Faulkenberry, MD, MPH
Department of Hematology and Medical Oncology
Emory University School of Medicine
36 Linden Ave NE, Atlanta, GA 30308, USA
Email: grey.faulkenberry@emory.edu
Phone: +1 404 778 1900

**Keywords (MeSH):** Speech Recognition Software; Medication Errors; Documentation; Mobile Applications; Benchmarking

**Word count (body, excluding abstract, references, tables, figures):** 3,996

---

"""


# Per JAMIA General Instructions: images are uploaded as separate files
# (paper/figure1.png|pdf, paper/figure2.png|pdf go into the portal's figure
# upload step), legends are provided at the end of the manuscript, and each
# alt text description goes in the main manuscript file directly under the
# figure legend, preceded by "Alt text:".

FIGURES_BLOCK = """# Figure Legends

**Figure 1. Word error rate and clinical term recall of the best on-device ASR model versus the best cloud API, per dataset.**
**a**, Word error rate (%, standard Whisper-normalized regime; see Methods) of the best on-device model and the best cloud API on each of the three clinical conversation datasets (OSCE respiratory interviews, n = 272; PriMock57 primary care, n = 57; Kazi et al. psychiatric, n = 71). Lower is better. **b**, Clinical term recall (%) of the best on-device model and the best cloud API on each dataset. Higher is better. Clinical term recall was computed as the proportion of UMLS medical concept spans in the reference transcript that were correctly transcribed.

Alt text: Two-panel grouped bar chart comparing the best on-device ASR model with the best cloud API on each of three clinical conversation datasets. Panel a shows word error rate (lower is better); panel b shows clinical term recall (higher is better). Datasets are OSCE respiratory interviews (n=272), PriMock57 primary care (n=57), and Kazi et al. psychiatric (n=71).

**Figure 2. ROVER hypothesis fusion yields small improvements over the best single on-device model, with the best cloud API as reference.**
For each dataset, the figure shows the word error rate of the best single on-device model (left point), the best two-model ROVER fusion pair (right point, connected by a grey line to display the fusion delta), and the best cloud API as a dashed horizontal reference. Fusion provided ≤ 0.82 percentage point improvements on all three datasets. On the OSCE respiratory interview dataset, the best fused on-device pair (parakeet-tdt-0.6b-v2 + sensevoice, 11.01%) did not surpass the best cloud API (Azure, 7.70%).

Alt text: Dot-and-line plot showing word error rate for the best single on-device model versus the best two-model ROVER fusion pair on each of three clinical conversation datasets, with the best cloud API drawn as a dashed horizontal reference line. Fusion produces small improvements over the best single on-device model on all three datasets but does not reach the best cloud API on any of them.
"""


def read_file(path: Path) -> str:
    return path.read_text()


def strip_first_heading(text: str) -> str:
    """Drop the first top-level '## Abstract'-style heading from a file fragment."""
    lines = text.splitlines()
    out = []
    dropped = False
    for line in lines:
        if not dropped and line.startswith('## '):
            dropped = True
            continue
        out.append(line)
    return '\n'.join(out)


def build_main_manuscript() -> str:
    print('[main] assembling main manuscript', flush=True)

    # Start with the full_manuscript.md and splice in title page + tables + refs + legends.
    manuscript = read_file(PAPER / 'full_manuscript.md')

    # Prepend title page.
    out = TITLE_PAGE + manuscript.lstrip()

    # Convert CSV tables to markdown and splice each into the main text at
    # its [[TABLEn]] marker (JAMIA: tables are placed where first cited).
    # Table 1 has 9 columns; the per-column "(n=272)" labels bloat the headers
    # and force four-line wraps in the rendered docx, so we strip them from the
    # column headers (keeping the OSCE asterisk) and state the sizes in the
    # caption instead. Only the header row contains "(n=...)".
    tbl1 = re.sub(r'\s*\(n=\d+\)', '', csv_to_md_table(PAPER / 'table1_wer.csv'))
    tbl2 = csv_to_md_table(PAPER / 'table2_ctr.csv')
    tbl3 = csv_to_md_table(PAPER / 'table3_cost.csv')

    table_blocks = {
        '[[TABLE1]]': f"""**Table 1. Word error rate (%) by model and dataset, under two normalization regimes.** Each dataset column has two sub-columns: **Std** uses the Whisper English text normalizer alone (the de facto leaderboard standard, used by the HuggingFace Open ASR Leaderboard and MLPerf), and **MP** is meaning-preserving WER, which layers on a normalization for clinical text (UK/US spelling, hyphenation, compound words, possessive eponyms, spaced acronyms, honorifics, dosing units, and conversational backchannels; full rules in Methods). Asterisked OSCE columns reflect the apostrophe-injected reference transcripts (see Methods). Parameter counts (millions) are approximate, drawn from model documentation and cross-checked against the deployed ONNX files; MedASR's count, which its distribution does not list, was measured directly from its ONNX model file. All on-device models were run as their published int8-quantized builds; in the Zipformer and Qwen3-ASR builds the upstream distribution leaves one component unquantized (the decoder and convolutional frontend, respectively). Cloud services do not disclose model sizes (Proprietary). Dataset sizes: OSCE n = 272, PriMock57 n = 57, psychiatric n = 71. ‡Google medical_conversation psychiatric WER is aggregated over 67 of 71 files; four of the six longest recordings had not completed within the 10-minute client-side wait and were not retried (see Methods).

{tbl1}""",
        '[[TABLE2]]': f"""**Table 2. Clinical term recall (%) by model and dataset, with bias-corrected and accelerated (BCa) bootstrap 95% confidence intervals (10,000 resamples; common random numbers within each dataset).** Clinical term recall is the percentage of UMLS medical concept spans in the reference transcript transcribed without error. All five cloud services, including AWS Transcribe Medical, were evaluated on the complete datasets (Google on 396 of 400 files; see Methods).

{tbl2}""",
        '[[TABLE3]]': f"""**Table 3. Cloud API cost for the full benchmark (~90 hours, 400 conversations).** All five services were evaluated on the full benchmark (Google on 396 of 400 files; see Methods). On-device inference incurs no per-encounter cost.

{tbl3}""",
    }
    for marker, block in table_blocks.items():
        if marker not in out:
            raise SystemExit(f'[error] {marker} marker missing from full_manuscript.md')
        out = out.replace(marker, block)

    # Build references section from references.md (numbered list at end of file).
    refs_text = read_file(PAPER / 'references.md')
    # Take only the "Formatted Reference List" block.
    marker = '## Formatted Reference List'
    if marker in refs_text:
        refs_block = refs_text.split(marker, 1)[1]
    else:
        refs_block = refs_text
    # Drop the "(NEJM AI style)" subtitle if present.
    refs_block = refs_block.replace(' (NEJM AI style)', '').strip()
    # Demote the section heading.
    refs_section = '\n# References\n\n' + refs_block + '\n'

    out = out + refs_section + '\n' + FIGURES_BLOCK

    return out


def build_supplementary() -> str:
    print('[supp] assembling supplementary tables', flush=True)
    supp_md = read_file(PAPER / 'supplementary_tables.md')

    # S1 table is already inline as markdown in supplementary_tables.md. Append S2-S4 as full tables.
    s2 = csv_to_md_table(PAPER / 'supplementary_table_S2_rover_pairs.csv')
    s3 = csv_to_md_table(PAPER / 'supplementary_table_S3_rover_triples.csv')
    s4 = csv_to_md_table(PAPER / 'supplementary_table_S4_drug_rates.csv')

    appendix = f"""

## Table S2 data

{s2}

## Table S3 data

{s3}

## Table S4 data

{s4}
"""
    return supp_md + appendix


def build_cover_letter() -> str:
    print('[cover] assembling cover letter', flush=True)
    return read_file(PAPER / 'cover_letter.md')


def post_process_docx(docx_path: Path) -> None:
    """Finalize a pandoc-generated docx for submission.

    Pandoc leaves the page geometry unset, double-spaces nothing, and gives
    every table column equal width, which makes the 9-column Table 1 wrap into
    four-line headers spanning six pages. This pass (1) sets explicit Letter
    geometry, (2) double-spaces body text per JAMIA, and (3) gives each table a
    fixed layout with content-aware column widths, single-spaced compact cells,
    a repeating header row, and no mid-row page breaks.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Twips
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    # 1. Explicit page geometry so column widths are deterministic (pandoc
    #    leaves pgSz unset, which makes rendered width locale-dependent).
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    text_width_twips = int(6.5 * 1440)

    # 2. Double-space body paragraphs only (doc.paragraphs excludes table cells).
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2.0

    # 3. Format every table: fixed layout, content-aware column widths,
    #    single-spaced compact cells, header repeat, no row splitting.
    for table in doc.tables:
        _format_table(table, text_width_twips, OxmlElement, qn, Pt, Twips)

    doc.save(str(docx_path))


def _table_profile(table) -> tuple:
    """Return (column-width fractions summing to ~1, font size pt) for a table.

    The four hand-tuned main/S4 tables are keyed by header signature (verified
    by rendering). Every other table (the S2/S3 round-robin dumps and any future
    table) gets content-aware widths: each column is sized in proportion to its
    longest cell, and the font is shrunk just enough that the widest column
    fits its share of the text width.
    """
    rows = table.rows
    headers = [c.text.strip() for c in rows[0].cells]
    ncol = len(headers)
    joined = ' '.join(headers).lower()
    h0 = headers[0].strip().lower()
    if 'params' in joined:                       # Table 1 (WER), 9 columns
        return [0.235, 0.11, 0.10] + [0.0925] * 6, 8.5
    if ncol == 5 and h0 == 'model':              # Table 2 (CTR)
        return [0.26, 0.12] + [0.2067] * 3, 9.0
    if ncol == 4 and h0 == 'service':            # Table 3 (cost)
        return [0.34, 0.22, 0.22, 0.22], 9.5
    if ncol == 8 and h0 == 'model':              # Table S4 (per-drug)
        return [0.18, 0.09] + [0.1217] * 6, 8.0

    # content-aware fallback (S2/S3 round-robin and any future table)
    maxlen = [1] * ncol
    for r in rows:
        for ci, c in enumerate(r.cells):
            if ci < ncol:
                maxlen[ci] = max(maxlen[ci], len(c.text.strip()))
    total = float(sum(maxlen))
    fractions = [m / total for m in maxlen]
    # shrink font until the widest column fits (~0.0075 in per char per pt in TNR)
    widest_share_in = max(fractions) * 6.5
    pt = widest_share_in / (max(maxlen) * 0.0075)
    font_pt = round(max(6.5, min(9.0, pt)), 1)
    return fractions, font_pt


def _format_table(table, text_width_twips, OxmlElement, qn, Pt, Twips) -> None:
    fractions, font_pt = _table_profile(table)
    widths_tw = [max(360, int(f * text_width_twips)) for f in fractions]

    tbl = table._tbl
    tblPr = tbl.tblPr

    # fixed layout (honor explicit widths instead of autofitting to content)
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    # total table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(sum(widths_tw)))
    table.autofit = False
    table.allow_autofit = False

    # rebuild the grid that defines column widths
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths_tw:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)

    for ri, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        # keep each row on one page (no mid-row page breaks)
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))
        # repeat the header row at the top of each page the table spans
        if ri == 0 and trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_tw):
                cell.width = Twips(widths_tw[ci])
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(font_pt)


def write_and_convert(md_text: str, stem: str, to_docx: bool = True,
                      double_space: bool = False) -> None:
    md_path = BUILD / f'{stem}.md'
    md_path.write_text(md_text)
    print(f'[write] {md_path} ({len(md_text)} chars)', flush=True)

    if to_docx:
        docx_path = BUILD / f'{stem}.docx'
        # Pandoc converts markdown to docx. Include resource-path so image
        # references resolve against the paper/ directory.
        extra = [
            '--resource-path', str(PAPER),
            '--standalone',
        ]
        try:
            pypandoc.convert_file(
                str(md_path),
                'docx',
                outputfile=str(docx_path),
                extra_args=extra,
            )
            print(f'[docx] {docx_path}', flush=True)
        except Exception as exc:  # noqa: BLE001
            print(f'[error] pandoc conversion failed for {stem}: {exc}', flush=True)
            raise

        if double_space:
            post_process_docx(docx_path)
            print(f'[formatted] {docx_path}', flush=True)


def main() -> None:
    main_md = build_main_manuscript()
    write_and_convert(main_md, 'stapes_manuscript', double_space=True)

    supp_md = build_supplementary()
    write_and_convert(supp_md, 'stapes_supplementary', double_space=True)

    cover_md = build_cover_letter()
    write_and_convert(cover_md, 'stapes_cover_letter')

    print('[done] submission package assembled in', BUILD, flush=True)


if __name__ == '__main__':
    sys.exit(main())
